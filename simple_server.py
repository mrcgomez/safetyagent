#!/usr/bin/env python3
"""
Simple SafetyAgent AI Server
A simplified version that works without complex dependencies
"""

import os
import sys
import json
import logging
from typing import Dict, List, Optional
from pathlib import Path

# FastAPI imports
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.responses import HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import uvicorn

# Document processing imports
import docx2txt
from docx import Document
import PyPDF2
import io

# OpenAI integration
import openai
from openai import OpenAI

# Load environment variables from .env file
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    # dotenv not available, continue without it
    pass

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Initialize FastAPI app
app = FastAPI(title="SafetyAgent AI", version="1.0.0")

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Mount static files
app.mount("/static", StaticFiles(directory="frontend"), name="static")

# Global knowledge base
knowledge_base: Dict[str, Dict] = {}
stats = {
    'total_documents': 0,
    'total_chunks': 0,
    'categories': []
}

class SimpleDocumentProcessor:
    """Simple document processor without complex dependencies"""
    
    def __init__(self):
        self.chunk_size = 1000
        self.chunk_overlap = 200
    
    def extract_text(self, file_path: str) -> str:
        """Extract text from various file formats"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            return self._extract_pdf_text(file_path)
        elif file_ext in ['.txt', '.md']:
            return self._extract_text_file(file_path)
        elif file_ext in ['.docx', '.doc']:
            return self._extract_word_text(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF files"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return f"[Error extracting PDF content: {e}]"
    
    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text files"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read()
        except Exception as e:
            logger.error(f"Error reading text file: {e}")
            return f"[Error reading text file: {e}]"
    
    def _extract_word_text(self, file_path: str) -> str:
        """Extract text from Word documents"""
        try:
            # Try docx2txt first
            text = docx2txt.process(file_path)
            if text and text.strip():
                return text
        except Exception as e:
            logger.warning(f"docx2txt failed: {e}")
        
        try:
            # Fallback to python-docx
            doc = Document(file_path)
            text_parts = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            return "\n".join(text_parts)
        except Exception as e:
            logger.error(f"Error extracting Word text: {e}")
            return f"[Error extracting Word content: {e}]"
    
    def chunk_text(self, text: str, metadata: Dict) -> List[Dict]:
        """Split text into chunks"""
        chunks = []
        words = text.split()
        
        for i in range(0, len(words), self.chunk_size - self.chunk_overlap):
            chunk_words = words[i:i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            
            if chunk_text.strip():
                chunk = {
                    'id': f"{metadata['doc_id']}_chunk_{len(chunks)}",
                    'text': chunk_text,
                    'metadata': {
                        **metadata,
                        'chunk_index': len(chunks),
                        'word_count': len(chunk_words)
                    }
                }
                chunks.append(chunk)
        
        return chunks

# Initialize document processor
processor = SimpleDocumentProcessor()

# Initialize OpenAI client
openai_client = None
try:
    # Try to get API key from environment variable
    api_key = os.getenv('OPENAI_API_KEY')
    if api_key:
        openai_client = OpenAI(api_key=api_key)
        logger.info("✅ OpenAI client initialized successfully")
    else:
        logger.warning("⚠️  OPENAI_API_KEY not found in environment variables")
        logger.info("💡 To enable AI features, set your OpenAI API key:")
        logger.info("   export OPENAI_API_KEY='your-api-key-here'")
except Exception as e:
    logger.error(f"❌ Error initializing OpenAI client: {e}")
    openai_client = None

@app.get("/", response_class=HTMLResponse)
async def read_root():
    """Serve the simple chat interface"""
    return HTMLResponse(content="""
    <!DOCTYPE html>
    <html>
    <head>
        <title>SafetyAgent AI Chat</title>
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
            }
            
            body {
                font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif;
                background: #f5f5f5;
                height: 100vh;
                display: flex;
                flex-direction: column;
            }
            
            .header {
                background: #001A70;
                color: white;
                padding: 15px 20px;
                text-align: center;
                box-shadow: 0 2px 4px rgba(0,0,0,0.1);
            }
            
            .header h1 {
                font-size: 1.5em;
                font-weight: 600;
            }
            
            .chat-container {
                flex: 1;
                display: flex;
                flex-direction: column;
                max-width: 800px;
                margin: 0 auto;
                width: 100%;
                background: white;
                box-shadow: 0 0 20px rgba(0,0,0,0.1);
            }
            
            .messages {
                flex: 1;
                overflow-y: auto;
                padding: 20px;
                display: flex;
                flex-direction: column;
                gap: 15px;
            }
            
            .message {
                max-width: 80%;
                padding: 12px 16px;
                border-radius: 18px;
                line-height: 1.4;
                word-wrap: break-word;
            }
            
            .user-message {
                background: #001A70;
                color: white;
                align-self: flex-end;
                border-bottom-right-radius: 4px;
            }
            
            .bot-message {
                background: #f1f3f4;
                color: #333;
                align-self: flex-start;
                border-bottom-left-radius: 4px;
            }
            
            .input-container {
                padding: 20px;
                border-top: 1px solid #e0e0e0;
                background: white;
            }
            
            .input-wrapper {
                display: flex;
                gap: 10px;
                align-items: center;
            }
            
            .message-input {
                flex: 1;
                padding: 12px 16px;
                border: 2px solid #e0e0e0;
                border-radius: 25px;
                font-size: 16px;
                outline: none;
                transition: border-color 0.2s;
            }
            
            .message-input:focus {
                border-color: #001A70;
            }
            
            .send-button {
                background: #001A70;
                color: white;
                border: none;
                border-radius: 50%;
                width: 45px;
                height: 45px;
                cursor: pointer;
                display: flex;
                align-items: center;
                justify-content: center;
                font-size: 18px;
                transition: background-color 0.2s;
            }
            
            .send-button:hover {
                background: #002A90;
            }
            
            .send-button:disabled {
                background: #ccc;
                cursor: not-allowed;
            }
            
            .typing-indicator {
                display: none;
                align-self: flex-start;
                background: #f1f3f4;
                color: #666;
                padding: 12px 16px;
                border-radius: 18px;
                border-bottom-left-radius: 4px;
                font-style: italic;
            }
            
            .welcome-message {
                text-align: center;
                color: #666;
                font-style: italic;
                margin: 20px 0;
            }
            
            .footer {
                background: #f8f9fa;
                border-top: 1px solid #e0e0e0;
                padding: 15px 20px;
                text-align: center;
                margin-top: auto;
            }
            
            .footer img {
                height: 30px;
                width: auto;
                opacity: 0.8;
            }
            
            @media (max-width: 600px) {
                .chat-container {
                    height: 100vh;
                }
                
                .message {
                    max-width: 90%;
                }
            }
        </style>
    </head>
    <body>
        <div class="header">
            <h1>🛡️ Safety Agent</h1>
        </div>
        
        <div class="chat-container">
            <div class="messages" id="messages">
                <div class="welcome-message">
                    Ask me anything about safety procedures, requirements, or policies!
                </div>
            </div>
            
            <div class="typing-indicator" id="typing">
                SafetyAgent is thinking...
            </div>
            
            <div class="input-container">
                <div class="input-wrapper">
                    <input 
                        type="text" 
                        id="messageInput" 
                        class="message-input" 
                        placeholder="Ask about safety procedures..."
                        autocomplete="off"
                    >
                    <button id="sendButton" class="send-button" onclick="sendMessage()">
                        ➤
                    </button>
                </div>
            </div>
        </div>
        
        <div class="footer">
            <img src="data:image/svg+xml;base64,PHN2ZyBpZD0iTGF5ZXJfMSIgZGF0YS1uYW1lPSJMYXllciAxIiB4bWxucz0iaHR0cDovL3d3dy53My5vcmcvMjAwMC9zdmciIHZpZXdCb3g9IjAgMCAyNTUuMTMgMTA4Ij48ZGVmcz48c3R5bGU+LmNscy0xe2ZpbGw6IzAwNDY5Mzt9PC9zdHlsZT48L2RlZnM+PHBhdGggY2xhc3M9ImNscy0xIiBkPSJNOS44LDcuNDZINjAuODl2OTMuMzVINDUuMzRjMi45Mi05LjgxLDIuNzYtMjEuODQuODItMzItLjgtNC4yMi0xLjY4LTguNDUtMi44My0xMi41OWwtLjA2LS4xLS4xNi40OWMtMy43MiwxMy42OC02LjI0LDMwLjExLTIsNDQuMjJIMzYuODVBNTcuODEsNTcuODEsMCwwLDEsMjguNDUsODFhNTcuNzIsNTcuNzIsMCwwLDEtLjE3LTEyLjU0YzEuMzQtMTEsNS0yMSw5LTMwLjg1LDEuODMtNC40Myw0LjE0LTguNDcsNi4xNS0xMi44OFYxMS42N0MzNS4yNSwyNiwyNy43MSw0MSwyMy41LDU3LjA3Yy0yLjYyLDEwLjA3LTMuNDksMjIuMy0uMTMsMzIuMjdBNTAuNTQsNTAuNTQsMCwwLDAsMjguOCwxMDAuODFIOS44WiIvPjxwYXRoIGNsYXNzPSJjbHMtMSIgZD0iTTkxLjEzLDY3LjA5VjY2YzAtMy41Ny0uMzYtNS0xLjczLTYuODItMS41OC0xLjk0LTQuNjMtMy4wNS04LjU1LTMuMDUtNi42NiwwLTEwLjUzLDMuNjEtMTAuNTMsOS44MWE3Ljg5LDcuODksMCwwLDAsMi4zNCw2LDMyLjYyLDMyLjYyLDAsMCwwLDYsMy44MmM1LDIuNTksNS45LDMuNDUsNS45LDYuMiwwLDIuNTUtMS4zMiw0LTMuNjEsNC0yLjc5LDAtMy41Ni0xLjI1LTMuNjYtNS45M0g3MGMwLDQuMTUuMjUsNS40OCwxLjIyLDcuMTcsMS40MiwyLjQ4LDQuODgsMy44Niw5LjM1LDMuODYsMy42NywwLDYuNzItLjg3LDguNDUtMi4zOXMyLjYtNC4yNywyLjYtNy40OGMwLTQuNDktMS42My02LjkyLTYuMjEtOS4zN2wtMy0xLjUyYy0zLjQ2LTEuODMtNC44My0zLjQxLTQuODMtNS42czEuMjItMy40NSwzLjIxLTMuNDVjMi4zOCwwLDMuMDksMS4zMSwzLjE1LDUuODVaIi8+PHBhdGggY2xhc3M9ImNscy0xIiBkPSJNMTAxLjc5LDc4di0yLjhjMC0zLjEuNzEtNC4zMiwyLjM4LTQuMzIsMi4yNSwwLDIuNiwxLjE3LDIuNiw4LjcxdjEuNTJsLS4xLDEuODNjLS4xLDMtLjY2LDQtMi40NCw0LTEuOTQsMC0yLjQ0LTEuMTEtMi40NC01LjM0Vjc4Wm0xMS41NCwwYTI5LDI5LDAsMCwwLS41NS02LjI1Yy0uNzctMy4yNi0zLjYyLTQuODQtOC44MS00Ljg0LTMsMC01LjE0LjYxLTYuNTYsMS44OC0xLjU4LDEuNDMtMi4yNCw0LTIuMjQsOC43bDAsLjcxLDAsMi45MWMwLDQuNjguNDEsNi4yLDEuODMsNy43M3MzLjY2LDIuMDksNy4wNywyLjA5YzMuMTIsMCw1LS40Niw2LjcyLTEuNzlzMi4zNC0zLjE1LDIuMzktNy40MloiLz48cGF0aCBjbGFzcz0iY2xzLTEiIGQ9Ik0xMzIuMDUsNzkuMjRjLS4yLDUuNC0uOTEsNi42Mi0zLjgyLDYuNjJhMi45NCwyLjk0LDAsMCwxLTMtMS44M2MtLjUxLTEuMzItLjc2LTQuNDgtLjc2LTEwLjEyLDAtMy41Mi4xLTYuMDYuMjUtOC4yNS4yLTIuOTUsMS4zOC00LjIzLDMuNzctNC4yM3MzLjQxLDEuMjgsMy40NSw0LjE4Yy4wNi40Ni4wNiwxLjEyLjExLDJoN1Y2NS42NmExNi4zMywxNi4zMywwLDAsMC0uNS00LjE4Yy0xLTMuMzEtNC43OS01LjM0LTEwLTUuMzQtMy40NiwwLTYuMjYuODctOC4xOSwyLjQ5LTIuNDksMi4xNS0zLjExLDUtMy4xMSwxNS4wNywwLDMuODcuMTYsNy40OC4zNiw5LjUxLjYxLDUuMjQsNC4xMiw3Ljg0LDEwLjYzLDcuODQsMy44MiwwLDctLjkyLDguNi0yLjU0czIuMTQtMy43MiwyLjE4LTkuMjdaIi8+PHBhdGggY2xhc3M9ImNscy0xIiBkPSJNMTUzLjI4LDc5LjM5VjgyLjNjMCwzLjM2LS43MSw0LjYyLTIuNTUsNC42Mi0xLjUyLDAtMi4wOC0uOS0yLjA4LTMuNCwwLTMuMzEuNzEtNC4xMywzLjU2LTQuMTNaTTE0OS4yMSw3NGMuMDYtMywuMzYtMy41NiwyLTMuNTZzMi4wOC43MSwyLjA4LDMuNTZ2MS45NGMtMS4wNywwLTEuNTgsMC0yLjEzLDAtNi41NywwLTguODYsMi4wOC04Ljg2LDguMTksMCw0LjY4LDEuNzQsNi44MSw1LjM1LDYuODEsMi4yOSwwLDMuODYtLjc2LDUuNjQtMi44OWwuMSwyLjU5aDYuNTFhMzEuMDYsMzEuMDYsMCwwLDEtLjI1LTMuNTZWNzUuODljMC00LjMzLS4yNS01LjU1LTEuNDItNi45M3MtMy42Mi0yLjEzLTYuODItMi4xM2ExMi40NiwxMi40NiwwLDAsMC02LjI2LDEuMzhjLTEuNzgsMS0yLjI0LDIuMTgtMi4yOSw1Ljc0WiIvPjxyZWN0IGNsYXNzPSJjbHMtMSIgeD0iMTYzLjkyIiB5PSI1Ni42NiIgd2lkdGg9IjYuNDYiIGhlaWdodD0iMzMuODgiLz48cGF0aCBjbGFzcz0iY2xzLTEiIGQ9Ik0xODUuODUsNzIuNTd2NS4zaDMuNjdWODUuNmExOS42NywxOS42NywwLDAsMS0zLjMxLjI1Yy0yLjM0LDAtMy40Ni0uNi00LTItLjQ2LTEuMjItLjU2LTMuMTUtLjU2LTkuODcsMC0xLjY5LjA1LTMuMDYuMS00LjI3bC4xMS0yLjc1Yy4xNS00LjIzLDEuMTctNS42NSw0LjA3LTUuNjUsMi40OSwwLDMuNTUsMS4yMiwzLjU1LDQuMTJsLjA2LDEuNzNoN2ExOCwxOCwwLDAsMC0uNDUtNS4xOWMtLjgyLTMuNzEtNC4zNy01LjY1LTEwLjQ4LTUuNjVhMTkuMDgsMTkuMDgsMCwwLDAtNC43NC41NkE4LjcsOC43LDAsMCwwLDE3NSw2My44OCwxMDEuNTIsMTAxLjUyLDAsMCwwLDE3NC40LDc1YzAsNC44OS4xMSw3LjIzLjU3LDkuMzcuODEsNC40Myw0LjUzLDYuNzIsMTAuNzMsNi43MmE0Mi40NCw0Mi40NCwwLDAsMCwxMC44My0xLjQ4di0xN1oiLz48cGF0aCBjbGFzcz0iY2xzLTEiIGQ9Ik0yMTEuNDcsNzkuMzl2Mi45YzAsMy4zNi0uNzIsNC42NC0yLjU1LDQuNjQtMS41MywwLTIuMDgtLjkyLTIuMDgtMy40MSwwLTMuMzEuNy00LjEzLDMuNTUtNC4xM1ptLTQuMDgtNS4zOWMuMDUtMywuMzYtMy41NSwyLTMuNTVzMi4wOS43LDIuMDksMy41NXYxLjkzYy0xLjA4LDAtMS41OSwwLTIuMTQsMC02LjU3LDAtOC44NiwyLjEtOC44Niw4LjIsMCw0LjY4LDEuNzQsNi44MSw1LjM1LDYuODEsMi4yOSwwLDMuODYtLjc2LDUuNjUtMi44OWwuMSwyLjU5aDYuNTFhMzEuNTgsMzEuNTgsMCwwLDEtLjI2LTMuNTdWNzUuODhjMC00LjMyLS4yNS01LjU0LTEuNDItNi45MnMtMy42Mi0yLjEzLTYuODEtMi4xM2ExMi42NCwxMi42NCwwLDAsMC02LjI3LDEuMzdjLTEuNzcsMS0yLjI0LDIuMi0yLjI5LDUuNzVaIi8+PHBhdGggY2xhc3M9ImNscy0xIiBkPSJNMjM4LjEzLDc0LjA1YzAtMi42OS0uMy0zLjc3LTEuNDItNS0xLjMyLTEuNDgtMy40NS0yLjE5LTYuNjEtMi4xOS01LjQ1LDAtOC41NSwyLjM5LTguNTUsNi41NiwwLDIuNjksMS4zMSw0LjQyLDQuNjksNi4xMWwxLjcyLjgxYzMuNTYsMS43OCw0LjI3LDIuNSw0LjI3LDQuMjNBMi4yNywyLjI3LDAsMCwxLDIyOS45NCw4N2MtMS41OCwwLTIuMTQtLjg3LTIuMTgtMy4xMWE5LjMzLDkuMzMsMCwwLDAsMC0xLjA3aC02LjMydi42MmMwLDUuMDksMi44MSw3LjM3LDguODYsNy4zNyw1LjIzLDAsOC40NC0yLjY0LDguNDQtNy4wN0E2LjI2LDYuMjYsMCwwLDAsMjM3LDc5LjM0LDE4LjM3LDE4LjM3LDAsMCwwLDIzMiw3Ni4xOWwtMS40Mi0uNzFjLTItMS0yLjc1LTEuODQtMi43NS0zYTEuODEsMS44MSwwLDAsMSwyLTJjMS4zNywwLDIsLjc2LDIsMi41OSwwLC4yNiwwLC42MS4xLDEuMDdaIi8+PHBhdGggY2xhc3M9ImNscy0xIiBkPSJNMjQzLjE3LDg2LjMySDI0MlY4OS45aC0uNjNWODYuMzJIMjQwLjF2LS42MWgzLjA3WiIvPjxwYXRoIGNsYXNzPSJjbHMtMSIgZD0iTTI0Ny41Miw4OS45aC0uNjJWNzguNjhjMC0uMjUsMC0uNjEsMC0uNzNsLS4yNy42Ni0xLDIuMzYtMS0yLjM2LS4yNy0uNjZjMCwuMTIsMCwuNDgsMCwuNzNWODkuOWgtLjZWODUuNzFoLjY1bC45NCwyLjIzYy4xLjI0LjIzLjU1LjI5LjcyLjA2LS4xNy4xOC0uNDcuMjgtLjcybC45Mi0yLjIzaC42OFoiLz48L3N2Zz4=" alt="SoCalGas">
        </div>
        
        <script>
            const messagesContainer = document.getElementById('messages');
            const messageInput = document.getElementById('messageInput');
            const sendButton = document.getElementById('sendButton');
            const typingIndicator = document.getElementById('typing');
            
            // Auto-focus input
            messageInput.focus();
            
            // Send message on Enter key
            messageInput.addEventListener('keypress', function(e) {
                if (e.key === 'Enter' && !e.shiftKey) {
                    e.preventDefault();
                    sendMessage();
                }
            });
            
            async function sendMessage() {
                const message = messageInput.value.trim();
                if (!message) return;
                
                // Add user message to chat
                addMessage(message, 'user');
                
                // Clear input and disable send button
                messageInput.value = '';
                sendButton.disabled = true;
                showTyping();
                
                try {
                    const response = await fetch('/api/chat', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json',
                        },
                        body: JSON.stringify({ query: message })
                    });
                    
                    const result = await response.json();
                    hideTyping();
                    addMessage(result.response, 'bot');
                    
                } catch (error) {
                    hideTyping();
                    addMessage('Sorry, I encountered an error. Please try again.', 'bot');
                    console.error('Error:', error);
                } finally {
                    sendButton.disabled = false;
                    messageInput.focus();
                }
            }
            
            function addMessage(text, sender) {
                const messageDiv = document.createElement('div');
                messageDiv.className = `message ${sender}-message`;
                messageDiv.textContent = text;
                
                messagesContainer.appendChild(messageDiv);
                messagesContainer.scrollTop = messagesContainer.scrollHeight;
            }
            
            function showTyping() {
                typingIndicator.style.display = 'block';
                messagesContainer.scrollTop = messagesContainer.scrollHeight;
            }
            
            function hideTyping() {
                typingIndicator.style.display = 'none';
            }
        </script>
    </body>
    </html>
    """)

@app.post("/api/upload")
async def upload_document(file: UploadFile = File(...), category: str = Form("general")):
    """Upload and process a document"""
    try:
        # Validate file type
        allowed_extensions = ['.pdf', '.txt', '.md', '.docx', '.doc']
        file_ext = os.path.splitext(file.filename)[1].lower()
        
        if file_ext not in allowed_extensions:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_ext} not supported. Allowed: {allowed_extensions}"
            )
        
        # Save uploaded file temporarily
        temp_path = f"temp_{file.filename}"
        with open(temp_path, "wb") as buffer:
            content = await file.read()
            buffer.write(content)
        
        try:
            # Extract text
            text = processor.extract_text(temp_path)
            
            # Create document metadata
            doc_id = f"doc_{len(knowledge_base)}"
            metadata = {
                'doc_id': doc_id,
                'filename': file.filename,
                'category': category,
                'file_size': len(content),
                'text_length': len(text)
            }
            
            # Chunk the text
            chunks = processor.chunk_text(text, metadata)
            
            # Store in knowledge base
            knowledge_base[doc_id] = {
                'metadata': metadata,
                'chunks': chunks,
                'full_text': text
            }
            
            # Update stats
            stats['total_documents'] += 1
            stats['total_chunks'] += len(chunks)
            if category not in stats['categories']:
                stats['categories'].append(category)
            
            logger.info(f"Processed document: {file.filename} -> {len(chunks)} chunks")
            
            return {
                "message": "Document uploaded and processed successfully",
                "doc_id": doc_id,
                "chunks_created": len(chunks),
                "filename": file.filename
            }
            
        finally:
            # Clean up temp file
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.post("/api/chat")
async def chat_with_agent(query: dict):
    """AI-powered chat with safety manual knowledge base"""
    try:
        question = query.get('query', '').strip()
        if not question:
            return {"response": "Please provide a question.", "confidence": 0.0, "sources": []}
        
        # Check if knowledge base is empty
        if not knowledge_base:
            return {
                "response": "I don't have access to the safety manual yet. Please restart the server to load the safety manual.",
                "confidence": 0.0,
                "sources": []
            }
        
        # Find relevant chunks using improved search
        relevant_chunks = find_relevant_chunks(question)
        
        if not relevant_chunks:
            return {
                "response": "I couldn't find specific information about that topic in the safety manual. Try asking about general safety requirements, PPE, emergency procedures, or incident reporting.",
                "confidence": 0.0,
                "sources": []
            }
        
        # If OpenAI is available, use it for intelligent responses
        if openai_client:
            return await generate_ai_response(question, relevant_chunks)
        else:
            # Fallback to simple response
            return generate_simple_response(relevant_chunks)
        
    except Exception as e:
        logger.error(f"Error in chat: {e}")
        return {
            "response": f"Sorry, I encountered an error: {str(e)}",
            "confidence": 0.0,
            "sources": []
        }

def find_relevant_chunks(question: str, max_chunks: int = 5):
    """Find relevant chunks from the knowledge base"""
    question_lower = question.lower()
    question_words = [word.strip('.,!?') for word in question_lower.split()]
    
    relevant_chunks = []
    for doc_id, doc_data in knowledge_base.items():
        for chunk in doc_data['chunks']:
            chunk_text_lower = chunk['text'].lower()
            
            # Calculate relevance score
            matches = sum(1 for word in question_words if word in chunk_text_lower)
            if matches > 0:
                # Boost score for exact phrase matches
                if any(phrase in chunk_text_lower for phrase in question_lower.split()):
                    matches += 2
                
                chunk['relevance_score'] = matches / len(question_words)
                relevant_chunks.append(chunk)
    
    # Sort by relevance and return top chunks
    relevant_chunks.sort(key=lambda x: x.get('relevance_score', 0), reverse=True)
    return relevant_chunks[:max_chunks]

async def generate_ai_response(question: str, relevant_chunks: list):
    """Generate AI response using OpenAI"""
    try:
        # Prepare context from relevant chunks
        context = "\n\n".join([chunk['text'] for chunk in relevant_chunks[:3]])
        
        # Create the prompt
        system_prompt = """You are SafetyAgent AI, a helpful assistant that answers questions about workplace safety based on the provided safety manual content. 

Your role:
- Answer safety-related questions clearly and accurately
- Use only information from the provided safety manual content
- Be helpful, professional, and safety-focused
- If the information isn't in the manual, say so clearly
- Provide specific, actionable guidance when possible

Always base your answers on the safety manual content provided."""

        user_prompt = f"""Based on the following safety manual content, please answer this question: "{question}"

Safety Manual Content:
{context}

Please provide a clear, helpful answer based on the safety manual information above."""

        # Call OpenAI API
        response = openai_client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=500,
            temperature=0.3
        )
        
        ai_response = response.choices[0].message.content.strip()
        
        # Prepare sources
        sources = [{
            'filename': chunk['metadata']['filename'],
            'chunk_id': chunk['id'],
            'relevance': chunk.get('relevance_score', 0.8)
        } for chunk in relevant_chunks[:3]]
        
        return {
            "response": ai_response,
            "confidence": 0.9,
            "sources": sources
        }
        
    except Exception as e:
        logger.error(f"Error generating AI response: {e}")
        # Fallback to simple response
        return generate_simple_response(relevant_chunks)

def generate_simple_response(relevant_chunks: list):
    """Generate simple response without AI"""
    top_chunks = relevant_chunks[:3]
    
    if len(top_chunks) == 1:
        response_text = f"Based on the safety manual:\n\n{top_chunks[0]['text'][:500]}..."
    else:
        response_text = "Based on the safety manual:\n\n"
        for i, chunk in enumerate(top_chunks, 1):
            response_text += f"{i}. {chunk['text'][:300]}...\n\n"
    
    sources = [{
        'filename': chunk['metadata']['filename'],
        'chunk_id': chunk['id'],
        'relevance': chunk.get('relevance_score', 0.8)
    } for chunk in top_chunks]
    
    return {
        "response": response_text,
        "confidence": min(0.8, max(0.5, top_chunks[0].get('relevance_score', 0.8))),
        "sources": sources
    }

@app.get("/api/stats")
async def get_stats():
    """Get knowledge base statistics"""
    return stats

@app.get("/health")
async def health_check():
    """Health check endpoint"""
    return {"status": "healthy", "documents": stats['total_documents'], "chunks": stats['total_chunks']}

def load_safety_manual():
    """Load the safety manual from environment variable or JSON file"""
    try:
        # First try to load from environment variable (more secure)
        safety_manual_json = os.getenv('SAFETY_MANUAL_JSON')
        safety_manual_compressed = os.getenv('SAFETY_MANUAL_COMPRESSED')
        
        if safety_manual_compressed:
            logger.info("Loading Safety Manual from compressed environment variable...")
            import gzip
            import base64
            # Decompress the data
            compressed_data = base64.b64decode(safety_manual_compressed)
            decompressed_json = gzip.decompress(compressed_data).decode('utf-8')
            safety_data = json.loads(decompressed_json)
        elif safety_manual_json:
            logger.info("Loading Safety Manual from environment variable...")
            safety_data = json.loads(safety_manual_json)
        else:
            # Fallback to JSON file
            json_path = "safety_manual.json"
            if not os.path.exists(json_path):
                logger.error(f"❌ Safety manual not found in environment variable or file: {json_path}")
                return
            
            logger.info("Loading Safety Manual from JSON file...")
            with open(json_path, 'r', encoding='utf-8') as f:
                safety_data = json.load(f)
            
        logger.info(f"Loaded JSON with {safety_data['metadata']['total_chunks']} chunks")
        
        # Create document metadata
        doc_id = "safety_manual"
        metadata = {
            'doc_id': doc_id,
            'filename': safety_data['metadata']['source_file'],
            'category': 'safety_manual',
            'file_size': len(safety_manual_json) if safety_manual_json else os.path.getsize(json_path),
            'text_length': safety_data['metadata']['total_characters'],
            'extraction_date': safety_data['metadata']['extraction_date'],
            'source': 'environment' if safety_manual_json else 'file'
        }
        
        # Convert JSON chunks to our format
        chunks = []
        for i, chunk_data in enumerate(safety_data['chunks']):
            chunk = {
                'id': chunk_data['id'],
                'text': chunk_data['content'],
                'metadata': {
                    **metadata,
                    'chunk_index': i,
                    'title': chunk_data['title'],
                    'word_count': chunk_data['word_count']
                }
            }
            chunks.append(chunk)
        
        # Store in knowledge base
        knowledge_base[doc_id] = {
            'metadata': metadata,
            'chunks': chunks,
            'full_text': safety_data['full_text'],
            'sections': safety_data['sections']
        }
        
        # Update stats
        stats['total_documents'] = 1
        stats['total_chunks'] = len(chunks)
        stats['categories'] = ['safety_manual']
        
        logger.info(f"✅ Safety manual loaded: {len(chunks)} chunks")
        logger.info(f"📊 Knowledge base now has {len(knowledge_base)} documents")
        return True
        
    except Exception as e:
        logger.error(f"❌ Error loading safety manual: {e}")
        import traceback
        traceback.print_exc()
        return False

if __name__ == "__main__":
    # Load safety manual on startup
    load_safety_manual()
    
    # Get port from environment variable (for Render deployment)
    port = int(os.getenv("PORT", 8000))
    
    # Start the server
    logger.info("🚀 Starting Safety Agent Server...")
    logger.info("📊 Knowledge Base Stats:")
    logger.info(f"   • Documents: {stats['total_documents']}")
    logger.info(f"   • Chunks: {stats['total_chunks']}")
    logger.info(f"   • Categories: {stats['categories']}")
    logger.info(f"   • Port: {port}")
    
    uvicorn.run(app, host="0.0.0.0", port=port, log_level="info")
